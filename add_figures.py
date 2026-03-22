#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
add_figures.py
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
R 분석 후 생성된 그래프 PNG 파일을 논문 docx에 자동으로 삽입합니다.

사용법:
    python add_figures.py                   # 기본 경로 사용
    python add_figures.py --figures-dir C:/Users/minta/ICBBS/r_results/figures
    python add_figures.py --dry-run         # 실제 삽입 없이 placeholder 목록만 확인
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

import os
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("[ERROR] python-docx 가 설치되어 있지 않습니다.\n"
             "        pip install python-docx --break-system-packages")

BASE_DIR = Path(__file__).parent
DOCX_IN  = BASE_DIR / "마이크로바이옴_CRC_예측_최종논문.docx"
DOCX_OUT = BASE_DIR / "마이크로바이옴_CRC_예측_최종논문_with_figures.docx"
FIG_DIR  = Path("C:/Users/minta/ICBBS/r_results/figures")

FIGURE_MAP = {
    "fig1":    ("fig1_roc_dev.png",           6.0),
    "fig2":    ("fig2_roc_ext.png",           6.0),
    "fig3":    ("fig3_perf_barplot.png",      6.5),
    "fig4":    ("fig4_feature_importance.png",6.0),
    "fig5":    ("fig5_generalizability.png",  6.0),
    "fig_cal": ("fig_cal.png",                6.5),
}

def find_placeholders(doc):
    found = []
    for i, para in enumerate(doc.paragraphs):
        if "##FIGURE_" in para.text:
            start = para.text.index("##FIGURE_") + len("##FIGURE_")
            end   = para.text.index("##", start)
            found.append((i, para.text[start:end], para))
    return found

def insert_image_at_paragraph(para, img_path, width_inches):
    for run in para.runs:
        run.text = ""
    p_elem = para._p
    for child in list(p_elem):
        if child.tag != qn("w:pPr"):
            p_elem.remove(child)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        for shd in pPr.findall(qn("w:shd")):
            pPr.remove(shd)

def process(docx_in, docx_out, fig_dir, dry_run=False):
    if not docx_in.exists():
        sys.exit(f"[ERROR] 논문 파일을 찾을 수 없습니다: {docx_in}")
    doc = Document(str(docx_in))
    placeholders = find_placeholders(doc)
    if not placeholders:
        print("[INFO] ##FIGURE_xxx## 마커가 없습니다.")
        return
    print(f"[INFO] {len(placeholders)}개의 그래프 placeholder 발견:")
    for _, mid, _ in placeholders:
        fname, w = FIGURE_MAP.get(mid, ("(unknown)", 6.0))
        fpath = fig_dir / fname
        status = "✓ 파일 존재" if fpath.exists() else "✗ 파일 없음"
        print(f"       ##FIGURE_{mid}## → {fname}  [{status}]")
    if dry_run:
        print("\n[DRY-RUN] 실제 삽입 없이 종료합니다.")
        return
    replaced = skipped = 0
    for _, marker_id, para in placeholders:
        if marker_id not in FIGURE_MAP:
            skipped += 1; continue
        fname, width = FIGURE_MAP[marker_id]
        fpath = fig_dir / fname
        if not fpath.exists():
            print(f"[SKIP]  {fname} 파일 없음 → placeholder 유지")
            skipped += 1; continue
        print(f"[INSERT] ##FIGURE_{marker_id}## ← {fname}")
        try:
            insert_image_at_paragraph(para, fpath, width)
            replaced += 1
        except Exception as e:
            print(f"[ERROR]  {marker_id} 삽입 실패: {e}")
            skipped += 1
    doc.save(str(docx_out))
    print(f"\n[완료] {replaced}개 그래프 삽입, {skipped}개 건너랐")
    print(f"[저장] {docx_out}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="R 분석 그래프를 논문 docx에 자동 삽입")
    parser.add_argument("--docx-in",     type=Path, default=DOCX_IN)
    parser.add_argument("--docx-out",    type=Path, default=DOCX_OUT)
    parser.add_argument("--figures-dir", type=Path, default=FIG_DIR)
    parser.add_argument("--dry-run",     action="store_true")
    args = parser.parse_args()
    print("=" * 60)
    print("  마이크로바이옴 CRC 논문 – 그래프 자동 삽입 도구")
    print("=" * 60)
    process(args.docx_in, args.docx_out, args.figures_dir, args.dry_run)
